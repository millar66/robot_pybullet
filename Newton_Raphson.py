#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Nov 28 10:18:45 2023

@author: lihui.liu
"""
import matplotlib.pyplot as plt
import numpy as np
import math

x = np.linspace(-10,10,400)
y = 3 * x**3 + 5 * x**2 + 7.25
y = x**4 + x**3 - x**2 -x
dy = 9 * x**2 + 10 * x
plt.plot(x,y)
plt.plot(x,dy)
plt.show()

x_0 = 10
x = x_0
for i in range(100):
    y = 3 * x**3 + 5 * x**2 + 7.25
    y = x**4 + x**3 - x**2 -x
    dy = 9 * x**2 + 10 * x
    x = x - (y/dy)
    if abs(y) < 0.0001:
        print(i)
        break
print(x)

# %%
import sympy
import pybullet as p
from docx import Document
import pickle

joint_positions=[0]*11
joint_pos_err=np.zeros((11, 3))
joint_orn_err=np.zeros((11, 3))
end_pos=np.zeros((1,3))
base_pos = np.append(end_pos,1).reshape(4,1)
numJoints = 8
alpha = [0, -np.pi/2, np.pi/2, np.pi/2, -np.pi/2, -np.pi/2, 0, 0, np.pi/2, np.pi/2, 0]
A = [0, 0, 0, -0.04951, 0.04951, 0, 0, -0.04050, 0, 0.01125, 0]
theta_rol = np.array([0., 0, 0, 0, 0, 0, 0, 0, 0, 0, 0])



theta_rol = np.array([0, 0.5, -0.5, 0.5, -0.5, -0.5, 0.1, 0.5, 0.5, 0.5, 0.5])



theta_rol[5] = theta_rol[5] + np.pi/2
theta_rol[8] = theta_rol[8] + np.pi/2
L = [0.27985, 0, 0.36330, 0, 0.36665, 0, 0, 0.55443-0.16, 0.16, 0, 0]

# theta_rol[7] = np.pi/3

Td = np.zeros((numJoints, 4, 4))
Tx = np.zeros((numJoints, 4, 4))
Ty = np.zeros((numJoints, 4, 4))
Tz = np.zeros((numJoints, 4, 4))
T_dot = []
T_simplify = []
# T_dot = np.zeros((numJoints, 4, 4))
T_joint = []
# T_joint.resize(12,4)
point_joint = np.zeros((numJoints+1, 3))
orn_cos_x = np.zeros(numJoints)
orn_sin_x = np.zeros(numJoints)
orn_cos_y = np.zeros(numJoints)
orn_sin_y = np.zeros(numJoints)
orn_cos_z = np.zeros(numJoints)
orn_sin_z = np.zeros(numJoints)

# L = sympy.symbols("L1:12")
# alpha = sympy.symbols("alpha1:12")
# A = sympy.symbols("a1:12")
# theta_rol = np.array([0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0])
theta = sympy.symbols("theta1:12")
theta1 = sympy.symbols("theta1")
theta2 = sympy.symbols("theta2")
theta3 = sympy.symbols("theta3")
theta4 = sympy.symbols("theta4")
theta5 = sympy.symbols("theta5")
theta6 = sympy.symbols("theta6")
theta7 = sympy.symbols("theta7")
theta8 = sympy.symbols("theta8")
theta9 = sympy.symbols("theta9")
theta10 = sympy.symbols("theta10")
theta11 = sympy.symbols("theta11")
T = []

for i in range(numJoints):
    if i != 6:
        T_NumJoint = sympy.Matrix([
            [sympy.cos(theta[i]),             -sympy.sin(theta[i]),               0,             A[i]],
            [np.cos(alpha[i])*sympy.sin(theta[i]), sympy.cos(theta[i])*np.cos(alpha[i]), -np.sin(alpha[i]), -np.sin(alpha[i])*L[i]],
            [sympy.sin(theta[i])*np.sin(alpha[i]), sympy.cos(theta[i])*np.sin(alpha[i]),  np.cos(alpha[i]),  np.cos(alpha[i])*L[i]],
            [0,                                0,                                 0,             1]])
    else:
        T_NumJoint = sympy.Matrix([
            [1.,  0., 0., 0.],
            [0.,  0., 1., (theta[i])],
            [0., -1., 0., 0.],
            [0.,  0., 0., 1.]])
    T.append(T_NumJoint)
    if i ==0 :
        T_dot.append(T_NumJoint)
    else:
        T_dot.append(T_dot[i-1] * T_NumJoint)
    T_simplify.append(T_dot[i].xreplace({n : round(n, 6) for n in T_dot[i].atoms(sympy.Number)}))
    T_joint.append(T_simplify[i] * base_pos)
    # np_joint = T_joint[i].subs([(theta1,0),(theta2,0),(theta3,3)])[0:3].T
    
    
    
    np_joint = sympy.lambdify(('theta1','theta2','theta3','theta4','theta5','theta6'
                               ,'theta7','theta8','theta9','theta10','theta11'),T_joint[i],"numpy")
    point_joint[i+1] = np_joint(theta_rol[0],theta_rol[1],theta_rol[2],theta_rol[3],theta_rol[4]
                                ,theta_rol[5],theta_rol[6],theta_rol[7],theta_rol[8],theta_rol[9]
                                ,theta_rol[10])[0:3].T
    p.addUserDebugLine(point_joint[i], point_joint[i+1], lineColorRGB=[0,0,1], lineWidth=5)


# with open('/home/lihui.liu/mnt/workspace/python/robot/robot_pybullet/expr.pickle', 'wb') as file:
#     pickle.dump(T_simplify, file)
# with open('expr.pickle', 'rb') as file:
#     expr = pickle.load(file)

# print(expr)
# file=Document()
# # doc = Document(path)
# file.add_heading("robot_T_1_8:")
# file.add_paragraph(point_joint.all())
# file.save('/home/lihui.liu/mnt/workspace/python/robot/robot_pybullet/T.docx')
# %%
p.removeAllUserDebugItems()






# %%

for i in range(numJoints):
    orn_cos_x[i] = np.cos(orn[i][0] + float(joint_orn_err[i][0])/1000)
    orn_sin_x[i] = np.sin(orn[i][0] + float(joint_orn_err[i][0])/1000)
    orn_cos_y[i] = np.cos(orn[i][1] + float(joint_orn_err[i][1])/1000)
    orn_sin_y[i] = np.sin(orn[i][1] + float(joint_orn_err[i][1])/1000)
    L_M = sympy.Matrix([[1.0,   0,   0,  pos[i][0] + joint_pos_err[i][0]/1000],
                        [  0, 1.0,   0,  pos[i][1] + joint_pos_err[i][1]/1000],
                        [  0,   0, 1.0,  pos[i][2] + joint_pos_err[i][2]/1000],
                        [  0,   0,   0,  1.0]])
    alpha_M = sympy.Matrix([[1, 0,             0,            0],
                            [0, orn_cos_x[i], -orn_sin_x[i], 0],
                            [0, orn_sin_x[i],  orn_cos_x[i], 0],
                            [0, 0,             0,            1]])
    Ty = sympy.Matrix([[orn_cos_y[i],    0, orn_sin_y[i], 0],
                       [0,               1, 0,            0],
                       [orn_sin_y[i]*-1, 0, orn_cos_y[i], 0],
                       [0,               0, 0,            1]])
    theta_M = sympy.Matrix([[sympy.cos(orn[i][1] + theta[i]), -sympy.sin(orn[i][1] + theta[i]), 0, 0],
                            [sympy.sin(orn[i][1] + theta[i]),  sympy.cos(orn[i][1] + theta[i]), 0, 0],
                            [0,                                0,                               1, 0],
                            [0,                                0,                               0, 1]])
    T.append(L_M * alpha_M * Ty * theta_M)
    T_dot = T_dot * T[i]
    T_joint.append(T_dot * base_pos)
    np_joint = sympy.lambdify(('theta1','theta2','theta3','theta4','theta5','theta6'
                               ,'theta7','theta8','theta9','theta10','theta11'),T_joint[i],"numpy")
    point_joint[i+1] = np_joint(theta_rol[0],theta_rol[1],theta_rol[2],theta_rol[3],theta_rol[4]
                                ,theta_rol[5],theta_rol[6],theta_rol[7],theta_rol[8],theta_rol[9]
                                ,theta_rol[10])[0:3].T
    p.addUserDebugLine(point_joint[i], point_joint[i+1], lineColorRGB=[0,0,1], lineWidth=5)

# T_joint[2][0].evalf(subs={theta1:theta[0],theta2:theta[1]},n=5)



x = sympy.symbols("x")
a = sympy.simplify(0.00001*sympy.sin(x)**2 + sympy.cos(x)**3)

b = a.xreplace({n : round(n, 4) for n in a.atoms(sympy.Number)})


# T_simplify = []
T_simplify = T_dot[7].xreplace({n : round(n, 6) for n in T_dot[7].atoms(sympy.Number)})


sympy.diff(T_simplify[7][0],theta8)
args = sympy.Matrix([theta1, theta2, theta3, theta4, theta5, theta6, theta7, theta8])
a = sympy.Matrix(T_simplify[7][0:4])
T_simplify[7].jacobian(args)

a.jacobian(args)






















